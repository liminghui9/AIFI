"""
报告导出模块
支持导出为Word和PDF格式
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak
from reportlab.lib import colors
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
import os
from typing import Dict, Optional
from datetime import datetime

class ExportGenerator:
    """报告导出生成器 - 专业版"""
    
    # 配色方案（参照UI设计图）
    COLOR_PRIMARY = RGBColor(41, 98, 255)      # 主蓝色 #2962FF
    COLOR_SECONDARY = RGBColor(66, 165, 245)   # 浅蓝色 #42A5F5
    COLOR_ACCENT = RGBColor(118, 75, 162)      # 紫色 #764BA2
    COLOR_SUCCESS = RGBColor(76, 175, 80)      # 绿色
    COLOR_WARNING = RGBColor(255, 152, 0)      # 橙色
    COLOR_DANGER = RGBColor(244, 67, 54)       # 红色
    COLOR_GRAY_DARK = RGBColor(66, 66, 66)     # 深灰
    COLOR_GRAY_LIGHT = RGBColor(245, 245, 245) # 浅灰
    
    def __init__(self):
        """初始化导出生成器"""
        self.doc = None
    
    def _add_header_footer(self, doc: Document, report_data: Dict):
        """添加页眉页脚"""
        section = doc.sections[0]
        
        # 页眉
        header = section.header
        header_para = header.paragraphs[0]
        header_para.text = f"企业财务分析报告 - {report_data['basic_info'].get('企业名称', '')}"
        header_para.style = doc.styles['Header']
        header_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        for run in header_para.runs:
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(128, 128, 128)
            run.font.name = '微软雅黑'
        
        # 页脚
        footer = section.footer
        footer_para = footer.paragraphs[0]
        footer_para.text = f"AIFI智能财报系统 | 生成时间: {report_data['generated_at']}"
        footer_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        for run in footer_para.runs:
            run.font.size = Pt(8)
            run.font.color.rgb = RGBColor(128, 128, 128)
            run.font.name = '微软雅黑'
    
    def _add_cover_page(self, doc: Document, report_data: Dict):
        """添加专业封面页"""
        # 封面标题
        title = doc.add_paragraph()
        title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        title_run = title.add_run('企业财务分析报告')
        title_run.font.size = Pt(36)
        title_run.font.bold = True
        title_run.font.color.rgb = self.COLOR_PRIMARY
        title_run.font.name = '微软雅黑'
        title_run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        
        # 副标题
        subtitle = doc.add_paragraph()
        subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        sub_run = subtitle.add_run('Financial Analysis Report')
        sub_run.font.size = Pt(14)
        sub_run.font.color.rgb = self.COLOR_SECONDARY
        sub_run.font.name = 'Arial'
        sub_run.font.italic = True
        
        # 添加空行
        for _ in range(2):
            doc.add_paragraph()
        
        # 企业名称
        company_name = doc.add_paragraph()
        company_name.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        name_run = company_name.add_run(report_data['basic_info'].get('企业名称', ''))
        name_run.font.size = Pt(24)
        name_run.font.color.rgb = self.COLOR_GRAY_DARK
        name_run.font.name = '微软雅黑'
        name_run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        
        # 添加空行
        for _ in range(5):
            doc.add_paragraph()
        
        # 生成信息表格（居中）
        info_table = doc.add_table(rows=5, cols=2)
        info_table.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        info_data = [
            ('报告类型', '企业财务分析报告'),
            ('企业名称', report_data['basic_info'].get('企业名称', '')),
            ('分析年度', f"{report_data['years'][0]}年" if report_data['years'] else '2023年'),
            ('生成时间', report_data['generated_at']),
            ('系统版本', 'AIFI v1.0')
        ]
        
        for i, (label, value) in enumerate(info_data):
            row = info_table.rows[i]
            row.cells[0].text = label
            row.cells[1].text = str(value)
            
            # 设置单元格样式
            for j, cell in enumerate(row.cells):
                cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                for run in cell.paragraphs[0].runs:
                    run.font.size = Pt(11)
                    run.font.name = '微软雅黑'
                    if j == 0:  # 标签列加粗
                        run.font.bold = True
                        run.font.color.rgb = self.COLOR_GRAY_DARK
                
                # 添加边框
                self._set_cell_border(cell)
        
        # 分页
        doc.add_page_break()
    
    def _set_cell_border(self, cell, **kwargs):
        """设置单元格边框"""
        tc = cell._element
        tcPr = tc.get_or_add_tcPr()
        
        # 创建边框元素
        tcBorders = OxmlElement('w:tcBorders')
        for edge in ('top', 'left', 'bottom', 'right'):
            edge_element = OxmlElement(f'w:{edge}')
            edge_element.set(qn('w:val'), 'single')
            edge_element.set(qn('w:sz'), '4')
            edge_element.set(qn('w:color'), 'CCCCCC')
            tcBorders.append(edge_element)
        
        tcPr.append(tcBorders)
    
    def _add_table_of_contents(self, doc: Document):
        """添加目录页"""
        # 目录标题
        toc_title = doc.add_heading('目  录', level=1)
        toc_title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        for run in toc_title.runs:
            run.font.color.rgb = self.COLOR_PRIMARY
            run.font.name = '微软雅黑'
        
        doc.add_paragraph()
        
        # 目录项
        toc_items = [
            ('一、企业基本信息', 1),
            ('二、总体风险评估', 2),
            ('三、数据可视化分析', 3),
            ('四、分维度风险分析', 4),
            ('    （1）盈利风险分析', 5),
            ('    （2）偿债风险分析', 6),
            ('    （3）运营风险分析', 7),
            ('    （4）现金流风险分析', 8),
            ('五、财务报表数据', 9),
            ('    （1）资产负债表', 10),
            ('    （2）利润表', 11),
            ('    （3）现金流量表', 12),
            ('免责声明', 13),
        ]
        
        for item, page in toc_items:
            p = doc.add_paragraph()
            p.add_run(item).font.size = Pt(12)
            p.add_run('\t' * 5 + f'{page}').font.size = Pt(12)
            p.paragraph_format.left_indent = Inches(0.5) if item.startswith('    ') else Inches(0)
        
        doc.add_page_break()
    
    def _add_styled_heading(self, doc: Document, text: str, level: int = 1):
        """添加带样式的标题"""
        heading = doc.add_heading(text, level=level)
        for run in heading.runs:
            run.font.color.rgb = self.COLOR_PRIMARY if level == 1 else self.COLOR_SECONDARY
            run.font.name = '微软雅黑'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        return heading
    
    def _create_styled_table(self, doc: Document, data: list, has_header: bool = True):
        """创建带样式的表格（优化版）"""
        table = doc.add_table(rows=len(data), cols=len(data[0]))
        
        # 设置表格宽度
        table.autofit = False
        table.allow_autofit = False
        
        for i, row_data in enumerate(data):
            row = table.rows[i]
            
            # 设置行高
            row.height = Cm(0.8)
            
            for j, cell_text in enumerate(row_data):
                cell = row.cells[j]
                cell.text = str(cell_text)
                
                # 设置单元格样式
                paragraph = cell.paragraphs[0]
                
                # 对齐方式：第一列左对齐，其他列居中
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER if j > 0 else WD_PARAGRAPH_ALIGNMENT.LEFT
                
                # 设置段落格式
                paragraph.paragraph_format.space_before = Pt(3)
                paragraph.paragraph_format.space_after = Pt(3)
                
                for run in paragraph.runs:
                    run.font.size = Pt(10)
                    run.font.name = '微软雅黑'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
                    
                    # 表头样式
                    if i == 0 and has_header:
                        run.font.bold = True
                        run.font.color.rgb = RGBColor(255, 255, 255)
                    # 第一列加粗
                    elif j == 0:
                        run.font.bold = True
                        run.font.color.rgb = self.COLOR_GRAY_DARK
                
                # 添加单元格边框
                self._set_cell_border(cell)
                
                # 表头背景色
                if i == 0 and has_header:
                    self._set_cell_background(cell, self.COLOR_PRIMARY)
                # 奇偶行背景色
                elif i % 2 == 0:
                    self._set_cell_background(cell, RGBColor(250, 250, 250))
        
        return table
    
    def _set_cell_background(self, cell, color: RGBColor):
        """设置单元格背景色"""
        shading_elm = OxmlElement('w:shd')
        # RGBColor对象直接包含RGB值，需要这样访问
        r, g, b = color
        shading_elm.set(qn('w:fill'), f'{r:02x}{g:02x}{b:02x}')
        cell._element.get_or_add_tcPr().append(shading_elm)
    
    def export_to_word(self, report_data: Dict, output_path: str) -> bool:
        """
        导出为Word格式（专业版）
        
        Args:
            report_data: 报告数据
            output_path: 输出文件路径
            
        Returns:
            bool: 是否成功
        """
        try:
            doc = Document()
            
            # 设置页面
            sections = doc.sections
            for section in sections:
                section.page_height = Inches(11.69)  # A4高度
                section.page_width = Inches(8.27)    # A4宽度
                section.top_margin = Cm(2.54)
                section.bottom_margin = Cm(2.54)
                section.left_margin = Cm(3.17)
                section.right_margin = Cm(3.17)
            
            # 1. 添加页眉页脚
            self._add_header_footer(doc, report_data)
            
            # 2. 添加封面
            self._add_cover_page(doc, report_data)
            
            # 3. 添加目录
            self._add_table_of_contents(doc)
            
            # 3. 企业基本信息
            self._add_styled_heading(doc, '一、企业基本信息', 1)
            basic_info = report_data['basic_info']
            
            # 准备表格数据
            info_data = [['项目', '内容']]
            for key, value in basic_info.items():
                info_data.append([key, str(value)])
            
            # 创建样式表格
            self._create_styled_table(doc, info_data, has_header=True)
            doc.add_paragraph()
            
            # 4. 总体风险评估
            self._add_styled_heading(doc, '二、总体风险评估', 1)
            
            # 添加评估内容框
            assessment_para = doc.add_paragraph()
            assessment_para.add_run(report_data['overall_assessment'])
            for run in assessment_para.runs:
                run.font.size = Pt(11)
                run.font.name = '微软雅黑'
            
            # 添加浅色背景
            assessment_para.paragraph_format.left_indent = Cm(0.5)
            assessment_para.paragraph_format.right_indent = Cm(0.5)
            assessment_para.paragraph_format.space_before = Pt(6)
            assessment_para.paragraph_format.space_after = Pt(6)
            
            doc.add_paragraph()
            
            # 4.1 数据可视化分析说明
            if report_data.get('charts'):
                self._add_styled_heading(doc, '三、数据可视化分析', 1)
                
                intro_para = doc.add_paragraph()
                intro_run = intro_para.add_run('本报告包含以下数据可视化图表：')
                intro_run.bold = True
                intro_run.font.size = Pt(11)
                intro_run.font.color.rgb = self.COLOR_GRAY_DARK
                intro_run.font.name = '微软雅黑'
                
                chart_descriptions = [
                    ('health_dashboard', '财务健康度仪表盘', '全面展示企业各维度财务健康状况'),
                    ('main_indicators', '主要财务指标对比图', '展示核心财务指标的年度对比'),
                    ('profitability', '盈利能力趋势图', '分析企业盈利能力的发展趋势'),
                    ('solvency', '偿债能力分析图', '评估企业短期和长期偿债能力'),
                    ('operational', '运营能力分析图', '衡量企业资产运营效率'),
                    ('cashflow', '现金流分析图', '分析企业现金流状况'),
                    ('risk_radar', '风险评估雷达图', '直观展示各维度风险等级')
                ]
                
                for chart_key, title, description in chart_descriptions:
                    if report_data['charts'].get(chart_key):
                        p = doc.add_paragraph()
                        bullet_run = p.add_run('●  ')
                        bullet_run.font.color.rgb = self.COLOR_PRIMARY
                        title_run = p.add_run(f'{title}：')
                        title_run.bold = True
                        title_run.font.size = Pt(10.5)
                        desc_run = p.add_run(description)
                        desc_run.font.size = Pt(10.5)
                        
                        for run in p.runs:
                            run.font.name = '微软雅黑'
                
                note_para = doc.add_paragraph()
                note_run = note_para.add_run('注：详细的交互式图表请查看在线报告页面进行分析。')
                note_run.italic = True
                note_run.font.size = Pt(9)
                note_run.font.color.rgb = RGBColor(128, 128, 128)
                note_run.font.name = '微软雅黑'
                
                doc.add_paragraph()
            
            # 5. 分维度风险分析
            self._add_styled_heading(doc, '四、分维度风险分析', 1)
            
            dimensions = ['盈利风险', '偿债风险', '运营风险', '现金流风险']
            dimension_icons = ['💰', '🛡️', '⚙️', '💧']
            years = report_data['years']
            
            for idx, dimension in enumerate(dimensions):
                # 维度标题
                dim_heading = self._add_styled_heading(doc, f'{dimension_icons[idx]} （{idx + 1}）{dimension}', 2)
                
                # 添加指标数据
                if years and dimension in report_data['indicators'].get(years[0], {}):
                    indicators = report_data['indicators'][years[0]][dimension]
                    
                    # 准备表格数据
                    table_data = [['指标名称', f'{years[0]}年']]
                    if len(years) >= 2:
                        table_data[0].append(f'{years[1]}年')
                        table_data[0].append('变化')
                    
                    # 填充数据
                    for indicator_name, value in indicators.items():
                        unit = self._get_indicator_unit(indicator_name)
                        row = [indicator_name]
                        
                        # 当前年度值
                        if value is not None:
                            row.append(f"{value:.2f}{unit}")
                        else:
                            row.append("数据缺失")
                        
                        # 上一年度值和变化
                        if len(years) >= 2:
                            prev_value = report_data['indicators'].get(years[1], {}).get(dimension, {}).get(indicator_name)
                            if prev_value is not None:
                                row.append(f"{prev_value:.2f}{unit}")
                                
                                # 计算变化
                                if value is not None:
                                    change = value - prev_value
                                    if change > 0:
                                        row.append(f"↑ {change:.2f}")
                                    elif change < 0:
                                        row.append(f"↓ {abs(change):.2f}")
                                    else:
                                        row.append("—")
                                else:
                                    row.append("—")
                            else:
                                row.append("数据缺失")
                                row.append("—")
                        
                        table_data.append(row)
                    
                    # 创建样式表格
                    self._create_styled_table(doc, table_data, has_header=True)
                
                # 添加AI分析框
                doc.add_paragraph()
                analysis_para = doc.add_paragraph()
                analysis_label = analysis_para.add_run('🤖 AI分析：')
                analysis_label.bold = True
                analysis_label.font.color.rgb = self.COLOR_PRIMARY
                analysis_label.font.size = Pt(10.5)
                
                analysis_text = analysis_para.add_run(report_data['dimension_analyses'].get(dimension, '暂无分析'))
                analysis_text.font.size = Pt(10.5)
                
                for run in analysis_para.runs:
                    run.font.name = '微软雅黑'
                
                # 添加背景
                analysis_para.paragraph_format.left_indent = Cm(0.5)
                analysis_para.paragraph_format.space_before = Pt(6)
                analysis_para.paragraph_format.space_after = Pt(6)
                
                doc.add_paragraph()
            
            # 6. 财务报表数据
            self._add_styled_heading(doc, '五、财务报表数据（精简版）', 1)
            
            financial_data = report_data['financial_data']
            tables_to_show = ['资产负债表', '利润表', '现金流量表']
            table_icons = ['📊', '💵', '💸']
            
            for idx, table_name in enumerate(tables_to_show):
                self._add_styled_heading(doc, f'{table_icons[idx]} （{idx + 1}）{table_name}', 2)
                
                if years:
                    year1_data = financial_data.get(years[0], {}).get(table_name, {})
                    
                    if year1_data:
                        # 准备表格数据
                        table_data = [['项目', f'{years[0]}年（万元）']]
                        if len(years) >= 2:
                            table_data[0].append(f'{years[1]}年（万元）')
                        
                        # 填充数据
                        for item_name, value in year1_data.items():
                            row = [item_name]
                            
                            if value is not None:
                                row.append(f"{value:,.2f}")
                            else:
                                row.append("数据缺失")
                            
                            if len(years) >= 2:
                                year2_value = financial_data.get(years[1], {}).get(table_name, {}).get(item_name)
                                if year2_value is not None:
                                    row.append(f"{year2_value:,.2f}")
                                else:
                                    row.append("数据缺失")
                            
                            table_data.append(row)
                        
                        # 创建样式表格
                        self._create_styled_table(doc, table_data, has_header=True)
                
                doc.add_paragraph()
            
            # 7. 添加免责声明
            doc.add_page_break()
            self._add_styled_heading(doc, '免责声明', 1)
            
            disclaimer = """本报告由AIFI智能财报系统自动生成，分析结果基于大语言模型和财务数据计算得出。

报告内容仅供参考，不构成审计、投资、法律或信用评级结论。

使用者应结合实际业务情况，谨慎判断和使用本报告信息。系统不对分析结果的准确性、完整性或稳定性作出保证。"""
            
            disclaimer_para = doc.add_paragraph(disclaimer)
            for run in disclaimer_para.runs:
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(128, 128, 128)
                run.font.name = '微软雅黑'
            
            disclaimer_para.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
            
            # 添加底部签名
            doc.add_paragraph()
            doc.add_paragraph()
            signature = doc.add_paragraph('—— AIFI 智能财报系统')
            signature.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
            for run in signature.runs:
                run.font.size = Pt(10)
                run.font.color.rgb = self.COLOR_PRIMARY
                run.font.italic = True
                run.font.name = '微软雅黑'
            
            # 保存文档
            doc.save(output_path)
            return True
            
        except Exception as e:
            print(f"导出Word失败: {str(e)}")
            import traceback
            traceback.print_exc()
            return False
    
    def export_to_pdf(self, report_data: Dict, output_path: str) -> bool:
        """
        导出为PDF格式（专业版）
        
        Args:
            report_data: 报告数据
            output_path: 输出文件路径
            
        Returns:
            bool: 是否成功
        """
        try:
            # 注册中文字体
            chinese_font = self._register_chinese_fonts()
            
            # 创建PDF文档
            doc = SimpleDocTemplate(
                output_path, 
                pagesize=A4,
                rightMargin=72,
                leftMargin=72,
                topMargin=72,
                bottomMargin=72
            )
            story = []
            styles = getSampleStyleSheet()
            
            # 自定义样式（使用中文字体）
            if chinese_font:
                # 封面标题样式
                cover_title_style = ParagraphStyle(
                    'CoverTitle',
                    parent=styles['Title'],
                    fontName=chinese_font,
                    fontSize=28,
                    alignment=1,  # 居中
                    textColor=colors.HexColor('#2962FF'),
                    spaceAfter=30,
                    spaceBefore=100
                )
                
                # 企业名称样式
                company_style = ParagraphStyle(
                    'CompanyName',
                    parent=styles['Title'],
                    fontName=chinese_font,
                    fontSize=18,
                    alignment=1,
                    textColor=colors.HexColor('#424242'),
                    spaceAfter=50
                )
                
                title_style = ParagraphStyle(
                    'CustomTitle',
                    parent=styles['Title'],
                    fontName=chinese_font,
                    fontSize=20,
                    alignment=1,
                    textColor=colors.HexColor('#2962FF'),
                    spaceAfter=12
                )
                
                heading1_style = ParagraphStyle(
                    'CustomHeading1',
                    parent=styles['Heading1'],
                    fontName=chinese_font,
                    fontSize=14,
                    textColor=colors.HexColor('#2962FF'),
                    spaceAfter=10,
                    spaceBefore=12
                )
                
                heading2_style = ParagraphStyle(
                    'CustomHeading2',
                    parent=styles['Heading2'],
                    fontName=chinese_font,
                    fontSize=12,
                    textColor=colors.HexColor('#42A5F5'),
                    spaceAfter=8,
                    spaceBefore=10
                )
                
                normal_style = ParagraphStyle(
                    'CustomNormal',
                    parent=styles['Normal'],
                    fontName=chinese_font,
                    fontSize=10,
                    leading=14
                )
                
                # 分析文本样式
                analysis_style = ParagraphStyle(
                    'Analysis',
                    parent=styles['Normal'],
                    fontName=chinese_font,
                    fontSize=10,
                    leading=14,
                    leftIndent=20,
                    rightIndent=20,
                    spaceAfter=6,
                    spaceBefore=6
                )
            else:
                # 如果没有中文字体，使用默认样式
                cover_title_style = styles['Title']
                company_style = styles['Title']
                title_style = styles['Title']
                heading1_style = styles['Heading1']
                heading2_style = styles['Heading2']
                normal_style = styles['Normal']
                analysis_style = styles['Normal']
            
            # 1. 封面页
            story.append(Paragraph('企业财务分析报告', cover_title_style))
            story.append(Spacer(1, 0.5*inch))
            
            # 企业名称
            story.append(Paragraph(report_data['basic_info'].get('企业名称', ''), company_style))
            story.append(Spacer(1, 1.5*inch))
            
            # 生成信息表格
            info_data = [
                ['报告类型', '企业财务分析报告'],
                ['生成时间', report_data['generated_at']],
                ['分析年度', f"{report_data['years'][0]}年" if report_data['years'] else '2023年'],
                ['系统版本', 'AIFI v1.0']
            ]
            
            info_table = Table(info_data, colWidths=[2*inch, 3*inch])
            table_font = chinese_font if chinese_font else 'Helvetica'
            info_table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, -1), colors.white),
                ('TEXTCOLOR', (0, 0), (-1, -1), colors.black),
                ('ALIGN', (0, 0), (0, -1), 'RIGHT'),
                ('ALIGN', (1, 0), (1, -1), 'LEFT'),
                ('FONTNAME', (0, 0), (-1, -1), table_font),
                ('FONTSIZE', (0, 0), (-1, -1), 10),
                ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
                ('TOPPADDING', (0, 0), (-1, -1), 8),
                ('BOTTOMPADDING', (0, 0), (-1, -1), 8),
            ]))
            story.append(info_table)
            
            # 分页
            story.append(PageBreak())
            
            # 2. 目录
            story.append(Paragraph('目  录', heading1_style))
            story.append(Spacer(1, 0.2*inch))
            
            toc_items = [
                '一、企业基本信息',
                '二、总体风险评估',
                '三、数据可视化分析',
                '四、分维度风险分析',
                '五、财务报表数据',
                '免责声明'
            ]
            
            for item in toc_items:
                story.append(Paragraph(f'  • {item}', normal_style))
            
            story.append(PageBreak())
            
            # 3. 企业基本信息
            story.append(Paragraph('一、企业基本信息', heading1_style))
            
            basic_info = report_data['basic_info']
            basic_data = [['项目', '内容']] + [[k, str(v)] for k, v in basic_info.items()]
            basic_table = Table(basic_data, colWidths=[2*inch, 4*inch])
            
            basic_table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#2962FF')),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                ('FONTNAME', (0, 0), (-1, -1), table_font),
                ('FONTSIZE', (0, 0), (-1, -1), 10),
                ('BOTTOMPADDING', (0, 0), (-1, 0), 8),
                ('TOPPADDING', (0, 0), (-1, -1), 6),
                ('BOTTOMPADDING', (1, 0), (-1, -1), 6),
                ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
            ]))
            story.append(basic_table)
            story.append(Spacer(1, 0.3*inch))
            
            # 4. 总体风险评估
            story.append(Paragraph('二、总体风险评估', heading1_style))
            story.append(Paragraph(report_data['overall_assessment'], analysis_style))
            story.append(Spacer(1, 0.3*inch))
            
            # 4.1 数据可视化分析（嵌入图片）
            if report_data.get('charts'):
                from reportlab.platypus import Image
                
                story.append(Paragraph('三、数据可视化分析', heading1_style))
                story.append(Spacer(1, 0.1*inch))
                
                # 图表列表及其PNG路径
                chart_list = [
                    ('health_dashboard_png', '财务健康度仪表盘'),
                    ('main_indicators_png', '主要财务指标对比'),
                    ('profitability_png', '盈利能力趋势'),
                    ('solvency_png', '偿债能力分析'),
                    ('operational_png', '运营能力分析'),
                    ('cashflow_png', '现金流分析'),
                    ('risk_radar_png', '风险评估雷达图')
                ]
                
                for chart_key, title in chart_list:
                    png_path = report_data['charts'].get(chart_key)
                    
                    if png_path and os.path.exists(png_path):
                        # 添加图表标题
                        story.append(Paragraph(f'<b>{title}</b>', heading2_style))
                        story.append(Spacer(1, 0.05*inch))
                        
                        try:
                            # 嵌入图片
                            img = Image(png_path, width=6*inch, height=3*inch)
                            story.append(img)
                            story.append(Spacer(1, 0.15*inch))
                        except Exception as e:
                            print(f"嵌入图表 {title} 失败: {str(e)}")
                            story.append(Paragraph('图表生成中...', normal_style))
                            story.append(Spacer(1, 0.1*inch))
                
                story.append(Spacer(1, 0.2*inch))
            
            # 5. 分维度风险分析
            story.append(Paragraph('四、分维度风险分析', heading1_style))
            
            dimensions = ['盈利风险', '偿债风险', '运营风险', '现金流风险']
            years = report_data['years']
            
            for dimension in dimensions:
                story.append(Paragraph(f'（{dimensions.index(dimension) + 1}）{dimension}', heading2_style))
                
                # 添加指标表格
                if years and dimension in report_data['indicators'].get(years[0], {}):
                    indicators = report_data['indicators'][years[0]][dimension]
                    
                    # 准备表格数据
                    table_data = [['指标名称', f'{years[0]}年']]
                    if len(years) >= 2:
                        table_data[0].append(f'{years[1]}年')
                    
                    for indicator_name, value in indicators.items():
                        unit = self._get_indicator_unit(indicator_name)
                        row = [indicator_name]
                        
                        if value is not None:
                            row.append(f"{value:.2f}{unit}")
                        else:
                            row.append("数据缺失")
                        
                        if len(years) >= 2:
                            prev_value = report_data['indicators'].get(years[1], {}).get(dimension, {}).get(indicator_name)
                            if prev_value is not None:
                                row.append(f"{prev_value:.2f}{unit}")
                            else:
                                row.append("数据缺失")
                        
                        table_data.append(row)
                    
                    # 创建表格
                    col_widths = [2.5*inch, 1.5*inch]
                    if len(years) >= 2:
                        col_widths.append(1.5*inch)
                    
                    indicator_table = Table(table_data, colWidths=col_widths)
                    
                    # 表格样式（使用中文字体）
                    indicator_table.setStyle(TableStyle([
                        ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                        ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                        ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                        ('FONTNAME', (0, 0), (-1, -1), table_font),
                        ('FONTSIZE', (0, 0), (-1, -1), 9),
                        ('BOTTOMPADDING', (0, 0), (-1, 0), 8),
                        ('TOPPADDING', (0, 0), (-1, -1), 4),
                        ('BOTTOMPADDING', (1, 0), (-1, -1), 4),
                        ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
                    ]))
                    story.append(indicator_table)
                
                # 添加分析
                story.append(Spacer(1, 0.1*inch))
                analysis_text = f"分析：{report_data['dimension_analyses'].get(dimension, '暂无分析')}"
                story.append(Paragraph(analysis_text, normal_style))
                story.append(Spacer(1, 0.2*inch))
            
            # 添加分页
            story.append(PageBreak())
            
            # 6. 免责声明
            story.append(Paragraph('免责声明', heading1_style))
            
            disclaimer_text = """本报告由AIFI智能财报系统自动生成，分析结果基于大语言模型和财务数据计算得出。<br/>
<br/>报告内容仅供参考，不构成审计、投资、法律或信用评级结论。<br/>
<br/>使用者应结合实际业务情况，谨慎判断和使用本报告信息。系统不对分析结果的准确性、完整性或稳定性作出保证。"""
            
            disclaimer_style = ParagraphStyle(
                'Disclaimer',
                parent=normal_style,
                fontSize=9,
                textColor=colors.grey,
                alignment=4  # 两端对齐
            )
            
            story.append(Paragraph(disclaimer_text, disclaimer_style))
            story.append(Spacer(1, 0.3*inch))
            
            # 签名
            signature_style = ParagraphStyle(
                'Signature',
                parent=normal_style,
                fontSize=10,
                textColor=colors.HexColor('#2962FF'),
                alignment=2,  # 右对齐
                fontName=table_font
            )
            story.append(Paragraph('<i>—— AIFI 智能财报系统</i>', signature_style))
            
            # 生成PDF
            doc.build(story)
            return True
            
        except Exception as e:
            print(f"导出PDF失败: {str(e)}")
            import traceback
            traceback.print_exc()
            return False
    
    def _register_chinese_fonts(self):
        """注册中文字体（如果可用）"""
        try:
            # 尝试注册常见的中文字体 - Windows
            font_paths = [
                ('SimSun', 'C:/Windows/Fonts/simsun.ttc'),      # 宋体
                ('SimHei', 'C:/Windows/Fonts/simhei.ttf'),      # 黑体
                ('Microsoft', 'C:/Windows/Fonts/msyh.ttc'),     # 微软雅黑
                ('SimSun', 'C:/Windows/Fonts/simsun.ttf'),      # 宋体备用
            ]
            
            for font_name, font_path in font_paths:
                if os.path.exists(font_path):
                    try:
                        pdfmetrics.registerFont(TTFont(font_name, font_path))
                        # 成功注册就返回字体名
                        return font_name
                    except:
                        continue
            
            # 如果都失败，返回None
            return None
        except:
            return None
    
    def _get_indicator_unit(self, indicator_name: str) -> str:
        """获取指标单位"""
        percentage_indicators = [
            '净利润率', '毛利率', '净资产收益率', '资产负债率'
        ]
        
        amount_indicators = [
            '经营性净现金流', '现金净增加额'
        ]
        
        if indicator_name in percentage_indicators:
            return '%'
        elif indicator_name in amount_indicators:
            return '万元'
        else:
            return ''

